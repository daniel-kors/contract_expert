import PyPDF2
from docx import Document
import re
from typing import Dict, List, Any, Optional


class FileProcessor:
    """Обработка файлов контрактов"""

    def extract_text(self, file_path: str) -> str:
        """Извлекает текст из файла с поддержкой разных регистров расширений"""
        try:
            file_path_lower = file_path.lower()

            if file_path_lower.endswith('.pdf'):
                return self._extract_from_pdf(file_path)
            elif file_path_lower.endswith('.docx'):
                return self._extract_from_docx(file_path)
            elif file_path_lower.endswith('.txt'):
                return self._extract_from_txt(file_path)
            else:
                return self._extract_with_fallback(file_path)
        except Exception as e:
            print(f"Error extracting text from {file_path}: {e}")
            return f"Ошибка чтения файла: {str(e)}"

    def _extract_from_pdf(self, file_path: str) -> str:
        """Извлекает текст из PDF файла"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)

                if reader.is_encrypted:
                    try:
                        reader.decrypt('')
                    except:
                        return "PDF файл защищен паролем"

                text = ""
                for page in reader.pages:
                    try:
                        text += page.extract_text() + "\n"
                    except:
                        continue

                return text if text.strip() else "Не удалось извлечь текст из PDF"

        except Exception as e:
            raise Exception(f"Ошибка чтения PDF: {str(e)}")

    def _extract_from_docx(self, file_path: str) -> str:
        """Извлекает текст из DOCX файла с улучшенной обработкой"""
        try:
            doc = Document(file_path)
            full_text = []

            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)

            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)

            return "\n".join(full_text)
        except Exception as e:
            raise Exception(f"Ошибка чтения DOCX: {str(e)}")

    def _extract_from_txt(self, file_path: str) -> str:
        """Извлекает текст из TXT файла"""
        try:
            encodings = ['utf-8', 'cp1251', 'windows-1251', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            raise Exception("Не удалось декодировать файл")
        except Exception as e:
            raise Exception(f"Ошибка чтения TXT: {str(e)}")

    def _extract_with_fallback(self, file_path: str) -> str:
        """Резервный метод извлечения текста"""
        try:
            return self._extract_from_txt(file_path)
        except:
            raise Exception("Неподдерживаемый формат файла")


class ContractValidator:
    """Улучшенная валидация контрактов"""

    def __init__(self):
        self.mandatory_clauses_44 = {
            "предмет контракта": ["предмет контракта", "предмет договора", "i. предмет", "1. предмет"],
            "цена контракта": ["цена контракта", "цена договора", "стоимость", "ii. цена", "2. цена"],
            "срок исполнения": ["срок исполнения", "срок поставки", "срок действия", "iii. порядок", "3. порядок"],
            "порядок оплаты": ["порядок оплаты", "условия оплаты", "расчеты", "оплата"],
            "ответственность сторон": ["ответственность сторон", "ответственность", "штраф", "пеня", "неустойка"],
        }

    def basic_validation(self, contract_text: str, law_type: str) -> Dict[str, Any]:
        """Улучшенная проверка обязательных условий"""
        errors = []
        warnings = []

        text_lower = contract_text.lower()

        # Проверяем наличие обязательных разделов
        for clause_name, patterns in self.mandatory_clauses_44.items():
            found = any(pattern in text_lower for pattern in patterns)
            if not found:
                errors.append({
                    'type': 'missing_clause',
                    'clause': clause_name,
                    'severity': 'critical',
                    'message': f'Отсутствует обязательный раздел: {clause_name}'
                })

        # ТОЧНЫЙ поиск цены контракта
        price_info = self._extract_price_info(contract_text)
        if not price_info.get('found'):
            errors.append({
                'type': 'missing_price',
                'severity': 'critical',
                'message': 'Не обнаружена цена контракта'
            })
        else:
            # Проверяем соответствие основанию заключения
            foundation_check = self._check_contract_foundation(contract_text, price_info['numeric_value'], law_type)
            if foundation_check:
                errors.append(foundation_check)

        return {
            'errors': errors,
            'warnings': warnings,
            'price_info': price_info
        }

    def _extract_price_info(self, text: str) -> Dict[str, Any]:
        """ТОЧНО извлекает информацию о цене контракта"""
        # Паттерны для поиска цены
        price_patterns = [
            r'цена\s+контракта\s+составляет\s+([^\n]{0,200}?)(\d{1,3}(?:\s?\d{3})*(?:[.,]\d{2})?)',
            r'стоимость[^\n]{0,100}?(\d{1,3}(?:\s?\d{3})*(?:[.,]\d{2})?)',
            r'(\d{1,3}(?:\s?\d{3})*(?:[.,]\d{2})?)\s*рубл',
        ]

        for pattern in price_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
            for match in matches:
                price_str = match.group(1) if len(match.groups()) > 1 else match.group(1)
                # Очищаем строку от пробелов и преобразуем в число
                price_clean = re.sub(r'[^\d,.]', '', price_str.replace(' ', ''))
                price_clean = price_clean.replace(',', '.')

                try:
                    numeric_value = float(price_clean)
                    return {
                        'found': True,
                        'text_value': price_str.strip(),
                        'numeric_value': numeric_value,
                        'context': match.group(0)[:100]
                    }
                except ValueError:
                    continue

        return {'found': False}

    def _check_contract_foundation(self, text: str, price: float, law_type: str) -> Optional[Dict]:
        """Проверяет соответствие основания заключения контракта его цене"""
        # Ищем упоминания статей 44-ФЗ
        article_93_pattern = r'ст\.\s*93|стать[ия]\s*93|п\.\s*25|пункт\s*25'
        article_93_matches = re.findall(article_93_pattern, text, re.IGNORECASE)

        if article_93_matches and price > 100000:
            return {
                'type': 'foundation_mismatch',
                'severity': 'critical',
                'message': f'Цена контракта ({price} руб.) превышает лимит для п. 25 ч. 1 ст. 93 44-ФЗ (100 000 руб.)'
            }
        elif article_93_matches and price <= 100000:
            return None  # Все корректно

        return None

    def compare_with_notice(self, contract_text: str, notice_text: str) -> Dict[str, Any]:
        """Сравнение контракта с извещением"""
        mismatches = []

        # Сравнение цены
        contract_price = self._extract_price_value(contract_text)
        notice_price = self._extract_price_value(notice_text)

        if contract_price and notice_price and contract_price != notice_price:
            mismatches.append({
                'parameter': 'price',
                'contract_value': contract_price,
                'notice_value': notice_price,
                'message': f'Цена контракта ({contract_price}) не соответствует цене в извещении ({notice_price})'
            })

        return {
            'mismatches': mismatches,
            'parameters_compared': ['price']
        }

    def _extract_price_value(self, text: str) -> Optional[str]:
        """Извлекает числовое значение цены"""
        price_info = self._extract_price_info(text)
        return price_info.get('text_value') if price_info.get('found') else None